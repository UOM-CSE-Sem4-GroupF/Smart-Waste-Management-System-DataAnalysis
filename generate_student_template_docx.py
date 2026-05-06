from docx import Document
from docx.shared import Pt

def md_to_docx(md_path: str, docx_path: str) -> None:
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code = False
    for raw in lines:
        line = raw.rstrip('\n')
        if not line:
            doc.add_paragraph('')
            continue

        if line.startswith('```'):
            in_code = not in_code
            if not in_code:
                doc.add_paragraph('')
            continue

        if in_code:
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            continue

        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('---'):
            doc.add_page_break()
        elif line.startswith('- '):
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
        elif line.startswith('+ '):
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
        elif line.startswith('Date |') and '---' in line:
            # table header line, start a simple table
            # read following line for table row
            # fallback: add as text
            doc.add_paragraph(line)
        else:
            doc.add_paragraph(line)

    doc.save(docx_path)


if __name__ == '__main__':
    md_to_docx('Student_Contribution_Template.md', 'Student_Contribution_Template.docx')
    print('Generated Student_Contribution_Template.docx')